import pinecone
from langchain_community.document_loaders.csv_loader import CSVLoader
from langchain_text_splitters import CharacterTextSplitter
from openai import AzureOpenAI
from pinecone import Pinecone, ServerlessSpec
from pinecone_text.sparse import BM25Encoder
from docx import Document
import pandas as pd
import warnings
import os
import nltk
from rank_bm25 import BM25Okapi
from sentence_transformers import SentenceTransformer
from openai_call import LLChatClient


# Download necessary NLTK data
nltk.download('punkt')

# Disable warnings
warnings.filterwarnings("ignore")
print("\033[31m!!!!!!ALL WARNINGS ARE DISABLED!!!!!!!!\033[0m")

# API Keys
openai_api_base_local = "https://eygenaistudio-openai-dev.openai.azure.com/"
openai_api_version_local = "2024-05-01-preview"
openai_api_key_local = "07362a79cd45421b99f245cda89c71fb"
pinecone_api_key = "pcsk_4NZkAC_NmmgNw869dExgqdZPesRhhLGDbQshb5Sg71K5vH5LHxAGcnTL7CxcPpHKAZwbU6"

# Initialize Pinecone & Azure OpenAI
pc = Pinecone(api_key=pinecone_api_key, ssl_verify=False)
client = AzureOpenAI(api_key=openai_api_key_local, api_version=openai_api_version_local, azure_endpoint=openai_api_base_local)

# Load sentence embedding model
embedding_model = SentenceTransformer("all-MiniLM-L6-v2")

# File path for document
file_path = "c:/Users/AT597VF/OneDrive - EY/Documents/Python Scripts/PRA/Annex_2.docx"

# Load and chunk data
def chunk_loader():
    print("Chunking started...")

    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File '{file_path}' not found.")
    
    doc = Document(file_path)
    text = "\n\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])

    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    text += "\n\n" + cell_text

    # Use a separator for splitting the text
    text_splitter = CharacterTextSplitter(separator="\n\n", chunk_size=500, chunk_overlap=50)
    chunks = text_splitter.split_text(text)  

    # Save chunks to CSV
    pd.DataFrame({'Chunk': chunks, 'Index': range(1, len(chunks) + 1)}).to_csv('extracted_tables.csv', index=False)
    print("Chunks saved to 'extracted_tables.csv'")

    return chunks

# Sparse Embedding (BM25)
def embed_sparse(chunks):
    print("Generating BM25 sparse embeddings...")
    bm25 = BM25Encoder.default()
    embedlist = []

    for i in chunks:
        doc_sparse_vector = bm25.encode_documents([i])

        if 'values' in doc_sparse_vector[0]:
            embedding_values = [float(x) for x in doc_sparse_vector[0]['values']]
        else:
            embedding_values = []

        embedlist.append(embedding_values)

    # Find the most common embedding dimension
    expected_dim = max(set(map(len, embedlist)), key=lambda d: sum(1 for e in embedlist if len(e) == d))
    print(f"Expected embedding dimension: {expected_dim}")

    # Standardize embeddings (truncate or pad)
    for i in range(len(embedlist)):
        if len(embedlist[i]) > expected_dim:
            embedlist[i] = embedlist[i][:expected_dim]  # Truncate
        elif len(embedlist[i]) < expected_dim:
            embedlist[i] += [0] * (expected_dim - len(embedlist[i]))  # Pad with zeros

    return embedlist


# Dense Embedding (Azure OpenAI)
def embed(chunks):
    print("Embedding chunks using OpenAI model...")
    embedlist = []  # List to store embeddings
    
    for i in chunks:
        response = client.embeddings.create(
            input=i, model="text-embedding-ada-002"
        )
        embedding = response.data[0].embedding
        
        # Ensure embeddings are 1536-dimensional
        if len(embedding) != 1536:
            raise ValueError(f"Unexpected embedding dimension: {len(embedding)}. Expected 1536.")

        embedlist.append(embedding)
    
    return embedlist



# Create Pinecone Index
def create_pinecone_index(index_name):
    pc = Pinecone(api_key=pinecone_api_key, ssl_verify=False)

    # Check if index exists
    if index_name in pc.list_indexes().names():
        print(f"Index '{index_name}' already exists.")

        # Get index metadata
        index_description = pc.describe_index(index_name)
        existing_dimension = index_description.dimension

        if existing_dimension != 1536:
            print(f"⚠️ Mismatched dimension! Existing index: {existing_dimension}, Expected: 1536")
            print("Deleting and recreating index with correct dimension...")
            delete_pinecone_index(index_name)

        else:
            print(f"✅ Index dimension matches (1536). No need to recreate.")
            return

    # Create new index with 1536 dimensions
    print(f"Creating index '{index_name}' with dimension 1536...")
    pc.create_index(
        name=index_name,
        dimension=1536,
        metric="cosine",  # Recommended for OpenAI embeddings
        spec=ServerlessSpec(cloud="aws", region="us-east-1"),
    )
    print(f"✅ Index '{index_name}' created successfully.")


# Delete Pinecone Index
def delete_pinecone_index(index_name="all"):
    if index_name == "all":
        indexes = pc.list_indexes().names()
        for index in indexes:
            pc.delete_index(index)
    else:
        pc.delete_index(index_name)

# Upsert Vectors into Pinecone
def upsert(embedlist):
    print("Upserting...")

    pc = Pinecone(api_key=pinecone_api_key, ssl_verify=False)
    index_name = "csai"

    # Check if index exists
    if index_name not in pc.list_indexes().names():
        raise ValueError(f"Index '{index_name}' does not exist. Please create it first.")

    index = pc.Index(index_name)  # Now safe to initialize

    local_vectors = []
    vector_id = 1
    for i in embedlist:
        vector_id_text = str(vector_id)
        float_values = [float(x) for x in i]
        local_vectors.append({"id": vector_id_text, "values": float_values})
        vector_id += 1

    batch_size = 300
    for i in range(0, len(local_vectors), batch_size):
        batch = local_vectors[i:i + batch_size]
        index.upsert(vectors=batch)

    print("\n✅ Upsert complete")


# Full pipeline function
def reupload():
    content = input("Do you want to update the knowledge base? Enter Y to confirm.\n")
    if content in ["Y", "y"]:
        chunks = chunk_loader()
        embedlist = embed(chunks)
        delete_pinecone_index("csai")
        # Ensure index exists before upserting
        index_name = "csai"
        create_pinecone_index(index_name)  

        upsert(embedlist)




# Run reupload
reupload()

# Example Query
query = "Exposures in default subject to a risk weight of 150%"
#results = hybrid_search(query)

# Print results
#for idx, res in enumerate(results):
#    print(f"{idx+1}. [Chunk ID: {res['Chunk_ID']}] Score: {res['Score']:.4f}\nText: {res['Text']}\n")


my_prompt = ChatPromptTemplate(
    input_variables=['content'],
    messages=[
        SystemMessage(content=sys_msg),
        #MessagesPlaceholder(variable_name='chat_history'), # Where the memory will be stored.
        HumanMessagePromptTemplate.from_template('{content}')
    ]
)
my_llm = AzureChatOpenAI(
    deployment_name="gpt-4",
    model_name="gpt-4",
    openai_api_key=openai_api_key_local,
    openai_api_type=openai_api_type_local,
    openai_api_version=openai_api_version_local,
    azure_endpoint=openai_api_base_local
)

client = LLMChatClient(llm, prompt, content)
response = client.get_response()
print(response)