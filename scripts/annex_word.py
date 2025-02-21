import pandas as pd
from docx import Document   

# def extract_tables_with_keywords(doc_path):
#     # Load the Word document
#     doc = Document(doc_path)

#     extracted_data = []  # List to store extracted rows with section keywords
#     current_keyword = "Unknown Section"  # Default section if no heading found
#     table_idx = 0  # Table counter
    
#     paragraphs = list(doc.paragraphs)  # Get all paragraphs
    
#     for i, para in enumerate(paragraphs):
#         text = para.text.strip()

#         # If paragraph is a potential section keyword (e.g., a heading)
#         if text and para.style.name.startswith("Heading"):  
#             current_keyword = text  # Update current keyword

#         # Check if a table follows this paragraph
#         if table_idx < len(doc.tables) and paragraphs[i+1].text.strip() == "":
#             table = doc.tables[table_idx]
#             headers = [cell.text.strip() for cell in table.rows[0].cells] if table.rows else []
#             rows = table.rows[1:]  # Skip header row

#     table_idx += 1  # Move to the next table

#     return extracted_data

def save_to_csv(extracted_data, output_csv):
    # Convert extracted data into a DataFrame
    df = pd.DataFrame(extracted_data, columns=["Section", "Row", "Instructions"])  # Adjust column names if needed
    df.to_csv(output_csv, index=False)  # Save to CSV

# Example usage
doc_path = "C:\\Users\\DP963UE\\OneDrive - EY\\Desktop\\HSBC_RC_POC\\Docs\\Annex 2 (Solvency).docx"
output_csv = "C:\\Users\\DP963UE\\OneDrive - EY\\Desktop\\HSBC_RC_POC\\extracted_tables.csv"

# Extract tables with keywords
# extracted_data = extract_tables_with_keywords(doc_path)
# print(extracted_data)
# Save to CSV
#save_to_csv(extracted_data, output_csv)

#print(f"Data successfully saved to {output_csv}")

def extract_all_tables(doc_path):
    # Load the Word document
    doc = Document(doc_path)

    all_tables_data = []  # List to store all table data
    index = 0  # Initialize index
    for table in doc.tables:
        #headers = [cell.text.strip() for cell in table.rows[0].cells] if table.rows else []
        
        for row in table.rows[1:]:  # Skip header row
            row_data = [cell.text.strip() for cell in row.cells]
            if len(row_data) >= 2 and any(row_data):  # Ensure there are at least two columns and row is not blank:  # Ensure there are at least two columns
                all_tables_data.append([str(row_data[0]), row_data[1],index])
                index += 1  # Increment index

    return all_tables_data

# Example usage
all_tables_data = extract_all_tables(doc_path)
df_all_tables = pd.DataFrame(all_tables_data, columns=["Reference", "Instruction", "Index"])


#print(df_all_tables)
save_to_csv(all_tables_data, output_csv)

